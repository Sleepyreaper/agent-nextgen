"""Document processing utilities for extracting text from various file formats."""

import base64
import logging
import os
import re
from typing import Any, Callable, Dict, List, Optional, Tuple

from docx import Document

logger = logging.getLogger(__name__)

# Minimum characters on a page before it's considered "image-based"
_IMAGE_PAGE_TEXT_THRESHOLD = 100

# Regex matching pagination-only footers like "Name - #1234\n7 of 13"
_PAGINATION_FOOTER_RE = re.compile(
    r'^\s*[\w\s,\'\-\.]+\s*-\s*#\d+\s*\n\s*\d+\s+of\s+\d+\s*$',
    re.IGNORECASE,
)


class DocumentProcessor:
    """Process and extract text from uploaded documents."""
    
    @staticmethod
    def extract_text_from_pdf(
        file_path: str,
        ocr_callback: Optional[Callable[[bytes, str], str]] = None,
    ) -> str:
        """Extract text from PDF file with page markers.
        
        Uses PyMuPDF (fitz) for text extraction.  When a page yields fewer
        than _IMAGE_PAGE_TEXT_THRESHOLD characters **and** contains an
        embedded image, the page is rendered to a PNG and passed to
        *ocr_callback* (if provided) so the caller can use an AI vision
        model to OCR it.
        
        Args:
            file_path: Path to the PDF file.
            ocr_callback: Optional ``fn(image_bytes, page_label) -> str``
                that performs OCR on a page image and returns extracted text.
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            # Fallback to pypdf if PyMuPDF is not installed
            return DocumentProcessor._extract_text_from_pdf_legacy(file_path)
        
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            text_parts: List[str] = []
            ocr_pages: List[int] = []
            
            # First pass: extract text and identify pages needing OCR
            pages_needing_ocr = []  # (page_idx, page_text, img_bytes, page_label)
            page_texts = {}  # page_idx → text
            
            for page_idx in range(total_pages):
                page = doc[page_idx]
                page_num = page_idx + 1
                
                # Extract tables first (preserves column structure for transcripts/grade tables)
                table_text = ""
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for tbl in tables.tables:
                            try:
                                df = tbl.to_pandas()
                                # Format as pipe-delimited with header row
                                header = " | ".join(str(c) for c in df.columns)
                                rows = []
                                for _, row in df.iterrows():
                                    rows.append(" | ".join(str(v) if v is not None else "" for v in row.values))
                                table_text += f"\n[TABLE]\n{header}\n" + "\n".join(rows) + "\n[/TABLE]\n"
                            except Exception:
                                pass
                except Exception:
                    pass  # find_tables() not available in older PyMuPDF versions
                
                # Get regular text content
                page_text = page.get_text().strip()
                
                # Append table text if tables were found (tables supplement regular text)
                if table_text.strip():
                    page_text = page_text + "\n" + table_text.strip() if page_text else table_text.strip()
                
                is_pagination_only = bool(
                    _PAGINATION_FOOTER_RE.match(page_text)
                ) if page_text else False
                effective_text_len = 0 if is_pagination_only else len(page_text)
                
                if effective_text_len < _IMAGE_PAGE_TEXT_THRESHOLD and ocr_callback:
                    images = page.get_images()
                    should_ocr = bool(images) or effective_text_len < 20 or is_pagination_only
                    if should_ocr:
                        try:
                            pix = page.get_pixmap(dpi=300)
                            img_bytes = pix.tobytes("png")
                            page_label = f"page {page_num} of {total_pages}"
                            pages_needing_ocr.append((page_idx, page_text, img_bytes, page_label, effective_text_len))
                            logger.info(f"🔍 Page {page_num}/{total_pages}: queued for OCR ({len(page_text)} chars, images={len(images) if images else 0})")
                        except Exception as e:
                            logger.warning(f"Failed to render page {page_num} for OCR: {e}")
                            page_texts[page_idx] = page_text
                    else:
                        page_texts[page_idx] = page_text
                elif effective_text_len < _IMAGE_PAGE_TEXT_THRESHOLD:
                    images = page.get_images()
                    if images or is_pagination_only:
                        print(
                            f"⚠️ Page {page_num}/{total_pages}: only {effective_text_len} effective chars "
                            f"(raw={len(page_text)}, pagination_only={is_pagination_only}, "
                            f"images={len(images) if images else 0}) — NO OCR callback available, text may be incomplete",
                            flush=True
                        )
                        logger.warning(
                            f"⚠️ Page {page_num}/{total_pages}: only {effective_text_len} effective chars "
                            f"(pagination_only={is_pagination_only}, images={len(images) if images else 0}) "
                            f"— no OCR callback, text may be incomplete"
                        )
                    elif effective_text_len < 50:
                        print(
                            f"⚠️ Page {page_num}/{total_pages}: sparse page ({effective_text_len} chars, 0 images, no OCR callback)",
                            flush=True
                        )
                
                if page_text:
                    page_texts[page_idx] = page_text
            
            # Parallel OCR pass: process all scanned pages concurrently (5-10x speedup)
            if pages_needing_ocr:
                logger.info(f"🚀 Parallel OCR: processing {len(pages_needing_ocr)} scanned pages...")
                from concurrent.futures import ThreadPoolExecutor, as_completed
                
                def _ocr_page(args):
                    pidx, ptext, img_bytes, plabel, eff_len = args
                    try:
                        ocr_text = ocr_callback(img_bytes, plabel)
                        if ocr_text and len(ocr_text.strip()) > eff_len:
                            return pidx, ocr_text.strip(), True
                        return pidx, ptext, False
                    except Exception as e:
                        logger.warning(f"OCR failed for {plabel}: {e}")
                        return pidx, ptext, False
                
                with ThreadPoolExecutor(max_workers=min(4, len(pages_needing_ocr))) as pool:
                    futures = {pool.submit(_ocr_page, args): args for args in pages_needing_ocr}
                    for future in as_completed(futures):
                        pidx, result_text, was_ocrd = future.result()
                        page_texts[pidx] = result_text
                        if was_ocrd:
                            ocr_pages.append(pidx + 1)
                            logger.info(f"✅ OCR extracted {len(result_text)} chars from page {pidx + 1}")
                
                logger.info(f"✅ Parallel OCR complete: {len(ocr_pages)} pages OCR'd")
            
            # Assemble final text in page order
            for page_idx in range(total_pages):
                pt = page_texts.get(page_idx, "")
                if pt:
                    text_parts.append(f"--- PAGE {page_idx + 1} of {total_pages} ---\n{pt}")
            
            doc.close()
            
            if ocr_pages:
                logger.info(f"📖 OCR was used on page(s): {ocr_pages}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from PDF with PyMuPDF: {e}")
            # Fallback to legacy pypdf
            return DocumentProcessor._extract_text_from_pdf_legacy(file_path)
    
    @staticmethod
    def _extract_text_from_pdf_legacy(file_path: str) -> str:
        """Legacy PDF extraction using pypdf (fallback)."""
        import pypdf
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                for i, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(f"--- PAGE {i} of {total_pages} ---\n{page_text}")
            return "\n\n".join(text)
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            return f"[Error reading PDF: {str(e)}]"
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document, including tables.
        
        Extracts both paragraph text and tabular data (common in
        transcripts, application forms, and recommendation letters
        formatted with Word tables).
        """
        try:
            doc = Document(file_path)
            parts: List[str] = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)
            
            # Extract tables (transcripts, grade tables, forms)
            for table_idx, table in enumerate(doc.tables):
                table_rows: List[str] = []
                # First row is likely the header — preserve it for column identity
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    # Deduplicate merged cells (Word repeats text for merged cells)
                    deduped: List[str] = []
                    for c in cells:
                        if not deduped or c != deduped[-1]:
                            deduped.append(c)
                    # Preserve empty cells as "—" so columns stay aligned
                    # (a missing credit-hours cell shouldn't shift Grade into Credits column)
                    deduped = [c if c else "—" for c in deduped]
                    row_text = " | ".join(deduped)
                    if row_text.replace("|", "").replace("—", "").strip():
                        table_rows.append(row_text)
                if table_rows:
                    parts.append(f"\n[Table {table_idx + 1}]")
                    parts.extend(table_rows)
            
            # Extract headers and footers
            for section in doc.sections:
                try:
                    header = section.header
                    if header and not header.is_linked_to_previous:
                        header_text = "\n".join(
                            p.text.strip() for p in header.paragraphs if p.text.strip()
                        )
                        if header_text:
                            parts.insert(0, header_text)
                except Exception:
                    pass
                try:
                    footer = section.footer
                    if footer and not footer.is_linked_to_previous:
                        footer_text = "\n".join(
                            p.text.strip() for p in footer.paragraphs if p.text.strip()
                        )
                        if footer_text:
                            parts.append(footer_text)
                except Exception:
                    pass
            
            return "\n\n".join(parts)
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            return f"[Error reading DOCX: {str(e)}]"
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                return f"[Error reading text file: {str(e)}]"
        except Exception as e:
            print(f"Error extracting text from TXT: {str(e)}")
            return f"[Error reading text file: {str(e)}]"
    
    @classmethod
    def process_document(
        cls,
        file_path: str,
        file_type: Optional[str] = None,
        ocr_callback: Optional[Callable[[bytes, str], str]] = None,
    ) -> Tuple[str, str]:
        """
        Process a document and extract its text.
        
        Args:
            file_path: Path to the document file
            file_type: Optional file type (will be inferred from extension if not provided)
            ocr_callback: Optional callback ``fn(image_bytes, page_label) -> str``
                for OCR of image-based PDF pages.
            
        Returns:
            Tuple of (extracted_text, detected_file_type)
        """
        # Determine file type
        if not file_type:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower().lstrip('.')
        
        # Extract text based on file type
        if file_type in ['pdf']:
            text = cls.extract_text_from_pdf(file_path, ocr_callback=ocr_callback)
        elif file_type in ['docx', 'doc']:
            text = cls.extract_text_from_docx(file_path)
        elif file_type in ['txt', 'text']:
            text = cls.extract_text_from_txt(file_path)
        else:
            text = f"[Unsupported file type: {file_type}]"
        
        return text, file_type
    
    @staticmethod
    def validate_file_type(filename: str) -> bool:
        """Check if the file type is supported (documents + video)."""
        allowed_extensions = {'pdf', 'docx', 'doc', 'txt', 'mp4'}
        _, ext = os.path.splitext(filename)
        return ext.lower().lstrip('.') in allowed_extensions

    @staticmethod
    def is_video_file(filename: str) -> bool:
        """Check if the file is a video file (requires Mirabel video analyzer)."""
        video_extensions = {'mp4'}
        _, ext = os.path.splitext(filename)
        return ext.lower().lstrip('.') in video_extensions
