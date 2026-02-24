"""Fairy Godmother - Document Generator Agent

Like the Fairy Godmother who transforms and creates beautiful outcomes,
this agent takes all the evaluation results and creates a polished, 
professional Word document summarizing the complete student evaluation.

This agent ALWAYS runs last after all other agents complete.
"""

import os
import json
import logging
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.agents.base_agent import BaseAgent

logger = logging.getLogger(__name__)


class FairyGodmotherDocumentGenerator(BaseAgent):
    """
    Fairy Godmother - Creates beautiful summary documents from evaluation results.
    
    Capabilities:
    - Generate professional Word documents
    - Include all agent evaluations
    - Format results beautifully
    - Store documents in student-specific folders
    - Provide download links
    """
    
    def __init__(self, db_connection=None, storage_manager=None):
        """
        Initialize Fairy Godmother Document Generator.
        
        Args:
            db_connection: Database connection for saving metadata
            storage_manager: Storage manager for file uploads
        """
        # Fairy Godmother doesn't need AI client - she creates documents directly
        super().__init__(name="Fairy Godmother", client=None)
        self.db = db_connection
        self.storage = storage_manager
        self.emoji = "🪄"
        self.description = "Creates beautiful evaluation summary documents"
    
    async def process(self, prompt: str) -> str:
        """Override base process - not used for document generation."""
        return "Fairy Godmother specializes in creating documents, not processing prompts."
    
    async def generate_evaluation_document(
        self,
        application: Dict[str, Any],
        agent_results: Dict[str, Any],
        student_id: str
    ) -> Dict[str, Any]:
        """
        Generate a comprehensive Word document with all evaluation results.
        
        Args:
            application: Student application data
            agent_results: Results from all agents (Tiana, Rapunzel, Moana, Mulan, Merlin, Aurora)
            student_id: Unique student identifier for storage
            
        Returns:
            Dict with document info and storage path
        """
        logger.info(f"🪄 Fairy Godmother creating evaluation document for {application.get('applicant_name', 'student')}")
        
        try:
            # Create Word document
            doc = Document()

            # Apply a consistent theme
            self._apply_document_theme(doc)
            
            # Add title and header
            self._add_header(doc, application)
            
            # Add executive summary
            self._add_executive_summary(doc, agent_results)
            
            # Add detailed agent results
            self._add_agent_results(doc, agent_results)
            
            # Add final recommendation
            self._add_final_recommendation(doc, agent_results)
            
            # Save document
            doc_path, doc_url = await self._save_document(doc, application, student_id)
            
            # Update database with document path
            if self.db and doc_path:
                await self._update_database(application, doc_path, doc_url)
            
            logger.info(f"✨ Fairy Godmother completed document: {doc_path}")
            
            return {
                'status': 'success',
                'document_path': doc_path,
                'document_url': doc_url,
                'student_id': student_id,
                'created_at': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Fairy Godmother error creating document: {e}", exc_info=True)
            return {
                'status': 'error',
                'error': str(e)
            }
    
    def _add_header(self, doc: Document, application: Dict[str, Any]):
        """Add document header with student info."""
        # Title
        title = doc.add_heading('Student Evaluation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph('NextGen Internship Admissions Review')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True

        date_line = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
        date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("-" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Student Information
        doc.add_heading('Applicant Information', level=1)
        
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light List Accent 1'
        
        info_data = [
            ('Name:', application.get('applicant_name') or application.get('ApplicantName', 'N/A')),
            ('Email:', application.get('email') or application.get('Email', 'N/A')),
            ('Application ID:', str(application.get('application_id') or application.get('ApplicationID', 'N/A'))),
            ('Evaluation Date:', datetime.now().strftime('%B %d, %Y'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
    
    def _add_executive_summary(self, doc: Document, agent_results: Dict[str, Any]):
        """Add executive summary from Aurora and Merlin."""
        doc.add_heading('Executive Summary', level=1)
        
        # Get Merlin's overall assessment
        merlin_data = agent_results.get('student_evaluator', {})
        aurora_data = agent_results.get('aurora', {})
        
        if merlin_data and merlin_data.get('status') == 'success':
            # Overall Score
            overall_score = merlin_data.get('overall_score', 'N/A')
            recommendation = merlin_data.get('recommendation', 'N/A')
            
            summary_para = doc.add_paragraph()
            summary_para.add_run('Overall Score: ').bold = True
            summary_para.add_run(f"{overall_score}/100\n")
            summary_para.add_run('Recommendation: ').bold = True
            summary_para.add_run(f"{recommendation}\n\n")
            
            # Rationale
            rationale = merlin_data.get('rationale', '')
            if rationale:
                summary_para.add_run('Assessment: ').bold = True
                summary_para.add_run(rationale)

            key_strengths = merlin_data.get('key_strengths') or []
            if key_strengths:
                doc.add_paragraph('Key Strengths:', style='Heading 3')
                for item in key_strengths:
                    doc.add_paragraph(f"- {item}")

            key_risks = merlin_data.get('key_risks') or []
            if key_risks:
                doc.add_paragraph('Considerations:', style='Heading 3')
                for item in key_risks:
                    doc.add_paragraph(f"- {item}")
        
        # Add Aurora's summary if available
        if aurora_data and isinstance(aurora_data, dict):
            merlin_summary = aurora_data.get('merlin_summary', {})
            if merlin_summary:
                doc.add_paragraph()
                doc.add_paragraph(f"⭐ {merlin_summary.get('summary_text', '')}")
        
        doc.add_paragraph()
    
    def _add_agent_results(self, doc: Document, agent_results: Dict[str, Any]):
        """Add detailed results from each agent."""
        doc.add_heading('Detailed Agent Evaluations', level=1)
        
        # Define agent display information
        agents_info = [
            ('application_reader', 'Tiana', '📋 Application Analysis'),
            ('grade_reader', 'Rapunzel', '📊 Academic Performance'),
            ('school_context', 'Moana', '🏫 School Context'),
            ('recommendation_reader', 'Mulan', '✉️ Recommendations'),
            ('student_evaluator', 'Merlin', '🔮 Final Evaluation')
        ]
        
        for agent_id, agent_name, section_title in agents_info:
            agent_data = agent_results.get(agent_id)
            if not agent_data:
                continue
            
            # Add section header
            doc.add_heading(section_title, level=2)
            
            if agent_data.get('status') == 'success':
                self._format_agent_output(doc, agent_name, agent_data)
            else:
                doc.add_paragraph(f"⚠️ {agent_name} evaluation not available")
            
            doc.add_paragraph()
    
    def _format_agent_output(self, doc: Document, agent_name: str, data: Dict[str, Any]):
        """Format individual agent output into the document."""
        # Remove status and other metadata fields
        display_data = {k: v for k, v in data.items() 
                       if k not in ['status', 'error', 'agent_name', 'timestamp']}
        
        for key, value in display_data.items():
            if isinstance(value, list):
                para = doc.add_paragraph()
                para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                for item in value:
                    doc.add_paragraph(f"- {item}")
            elif isinstance(value, dict):
                # Format complex data as JSON
                para = doc.add_paragraph()
                para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                try:
                    if any(style.name == 'Code' for style in doc.styles):
                        doc.add_paragraph(json.dumps(value, indent=2), style='Code')
                    else:
                        doc.add_paragraph(json.dumps(value, indent=2))
                except Exception:
                    doc.add_paragraph(json.dumps(value, indent=2))
            elif value is not None:
                para = doc.add_paragraph()
                para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                para.add_run(str(value))

    def _apply_document_theme(self, doc: Document):
        """Apply consistent fonts and spacing for a polished report."""
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.font.color.rgb = RGBColor(35, 46, 60)
        except Exception:
            pass

        for style_name, size, color in [
            ('Heading 1', 16, RGBColor(11, 31, 54)),
            ('Heading 2', 13, RGBColor(11, 31, 54)),
            ('Heading 3', 11, RGBColor(73, 87, 106))
        ]:
            try:
                style = doc.styles[style_name]
                style.font.name = 'Cambria'
                style.font.size = Pt(size)
                style.font.color.rgb = color
            except Exception:
                continue
    
    def _add_final_recommendation(self, doc: Document, agent_results: Dict[str, Any]):
        """Add final recommendation section."""
        doc.add_page_break()
        doc.add_heading('Final Recommendation', level=1)
        
        merlin_data = agent_results.get('student_evaluator', {})
        if merlin_data and merlin_data.get('status') == 'success':
            recommendation = merlin_data.get('recommendation', 'No recommendation available')
            confidence = merlin_data.get('confidence', 'Unknown')
            
            rec_para = doc.add_paragraph()
            rec_para.add_run('Decision: ').bold = True
            rec_para.add_run(f"{recommendation}\n")
            rec_para.add_run('Confidence: ').bold = True
            rec_para.add_run(f"{confidence}\n\n")
            
            # Add any additional notes
            notes = merlin_data.get('notes', '')
            if notes:
                rec_para.add_run('Additional Notes: ').bold = True
                rec_para.add_run(notes)
        
        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run('Generated by NextGen AI Evaluation System').italic = True
        footer_para.add_run(f"\n{datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    async def _save_document(
        self, 
        doc: Document, 
        application: Dict[str, Any],
        student_id: str
    ) -> tuple[str, Optional[str]]:
        """
        Save document to storage (local backup + Azure Storage).
        
        Returns:
            Tuple of (local_path, storage_url)
        """
        # Create filename
        applicant_name = application.get('applicant_name') or application.get('ApplicantName', 'student')
        safe_name = applicant_name.replace(' ', '_').replace('/', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"evaluation_{safe_name}_{timestamp}.docx"
        
        # Determine application type for proper container
        is_training = application.get('is_training_example', False)
        is_test = application.get('is_test_data', False)
        
        if is_training:
            application_type = 'training'
        elif is_test:
            application_type = 'test'
        else:
            application_type = '2026'
        
        # Create student-specific folder (local backup)
        base_dir = Path(os.environ.get('STUDENT_DOCUMENTS_DIR', str(Path(__file__).resolve().parent.parent.parent / 'student_documents')))
        student_dir = base_dir / application_type / student_id
        student_dir.mkdir(parents=True, exist_ok=True)
        
        # Save locally (backup)
        local_path = student_dir / filename
        doc.save(str(local_path))
        logger.info(f"Saved local backup: {local_path}")
        
        storage_url = None
        
        # Upload to Azure Storage (primary storage)
        if self.storage and hasattr(self.storage, 'client') and self.storage.client:
            try:
                with open(local_path, 'rb') as f:
                    file_content = f.read()
                
                upload_result = self.storage.upload_file(
                    file_content=file_content,
                    filename=filename,
                    student_id=student_id,
                    application_type=application_type
                )
                
                if upload_result.get('success'):
                    storage_url = upload_result.get('blob_url')
                    logger.info(f"✨ Document uploaded to Azure Storage: {storage_url}")
                    logger.info(f"   Container: {upload_result.get('container')}")
                    logger.info(f"   Path: {upload_result.get('blob_path')}")
                else:
                    logger.warning(f"Upload failed: {upload_result.get('error')}")
            except Exception as e:
                logger.warning(f"Could not upload to Azure Storage: {e}")
        else:
            logger.info("Azure Storage not configured, using local storage only")
        
        return str(local_path), storage_url
    
    async def _update_database(
        self, 
        application: Dict[str, Any],
        doc_path: str,
        doc_url: Optional[str]
    ):
        """Update database with document information."""
        application_id = application.get('application_id') or application.get('ApplicationID')
        
        if not application_id:
            logger.warning("No application ID to update database")
            return
        
        try:
            # Update application record with document path
            self.db.execute_non_query(
                """
                UPDATE applications 
                SET evaluation_document_path = %s,
                    evaluation_document_url = %s,
                    document_generated_at = NOW()
                WHERE application_id = %s
                """,
                (doc_path, doc_url, application_id)
            )
            logger.info(f"Updated database with document path for application {application_id}")
        except Exception as e:
            logger.error(f"Error updating database with document info: {e}", exc_info=True)
