"""Test document extraction improvements against real sample PDFs.

Tests the pagination-footer detection, OCR triggering logic,
DOCX table extraction, and Belle section adjacency fill.
"""

import os
import re
import sys

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.document_processor import DocumentProcessor, _PAGINATION_FOOTER_RE


# ── Pagination footer regex tests ──

def test_pagination_footer_regex():
    """Test that pagination-only footers are correctly detected."""
    positives = [
        "Thai, Brandon - #1807\n7 of 13",
        "Thomas, Dionna - #1370\n6 of 9",
        "Vinson, Kennedy - #1742\n4 of 7",
        "Whiskey, Micah - #1701\n6 of 9",
        "Uchenna, Victoria - #1851\n6 of 9",
        "Venzen, Jaide - #1852\n5 of 7",
        "Vides Angel, Jennifer - #1754\n5 of 7",
        "Williams, Kerrenton - #1497\n3 of 6",
        "Tapia Solis, Michel - #1846\n7 of 11",
        "Ablante, Brooke - #1645\n6 of 9",
        "Augustus, Chloe' - #1662\n5 of 10",
        "Brooks, Sy'ria - #1849\n4 of 7",
        "Kearse- Moorer, Eden - #1561\n4 of 8",
        "Aguilar Lopez, Jennifer - #1625\n3 of 5",
        "Barcenas-Lara, Abigail - #1744\n5 of 6",
        "Gonzalez-Langhorne, Olivia - #1600\n4 of 7",
        "Heywood-Peltier, Sierra - #1776\n3 of 5",
        "Velasquez-Lopez, Briseida - #1861\n5 of 9",
        "Venkadeshmeena, Sanjay - #1391\n5 of 9",
    ]
    negatives = [
        # Real content - should NOT match
        "Dear Emory Next Gen Program Selection Committee,",
        "McClure Health Science High School\nTranscript, Page 2\nStudent Name: Thai, Brandon",
        "Application Summary\n\nCompetition Details",
        "My name is Victoria Uchenna, and I am a dedicated magnet student",
        "District Name: Fulton County\nNorthview High School Transcript",
        "To Whom It May Concern:\nIt is my pleasure to strongly recommend",
        "Table of Contents\n\nThornton, Tessa - #1604 - Next Gen\n1\nessay_video",
        # Too much content after the footer pattern
        "Thai, Brandon - #1807\n7 of 13\nSome additional content here",
        # Empty
        "",
    ]

    print("=== Pagination Footer Regex Tests ===")
    passed = 0
    failed = 0

    for text in positives:
        match = bool(_PAGINATION_FOOTER_RE.match(text))
        if match:
            passed += 1
        else:
            failed += 1
            print(f"  FAIL (should match): {repr(text)}")

    for text in negatives:
        match = bool(_PAGINATION_FOOTER_RE.match(text))
        if not match:
            passed += 1
        else:
            failed += 1
            print(f"  FAIL (should NOT match): {repr(text)}")

    print(f"  {passed} passed, {failed} failed out of {len(positives) + len(negatives)}")
    return failed == 0


# ── PDF extraction tests (using sample PDFs if available) ──

def test_pdf_extraction_stats():
    """Verify extraction from sample PDFs identifies sparse/image pages correctly."""
    samples_dir = '/Users/sleepy/Downloads/OneDrive_1_3-2-2026'
    if not os.path.isdir(samples_dir):
        print("=== Skipping PDF extraction tests (samples not found) ===")
        return True

    test_cases = [
        # (filename, page_num (1-based), expected: 'sparse' if <50 chars without OCR)
        ("Thai, Brandon.pdf", 8, "sparse"),   # "Thai, Brandon - #1807\n7 of 13" + image
        ("Thai, Brandon.pdf", 9, "sparse"),   # same pattern
        ("Thai, Brandon.pdf", 10, "sparse"),  # same pattern
        ("Thomas, Dionna.pdf", 7, "sparse"),  # scanned rec letter
        ("Thomas, Dionna.pdf", 8, "sparse"),  # scanned rec letter
        ("Thomas, Dionna.pdf", 9, "sparse"),  # scanned rec letter
        ("Thai, Brandon.pdf", 3, "content"),  # real application text
        ("Thai, Brandon.pdf", 11, "content"), # transcript page with real text
    ]

    print("=== PDF Sparse Page Detection ===")
    import fitz
    passed = 0
    failed = 0

    for filename, page_num, expected in test_cases:
        path = os.path.join(samples_dir, filename)
        if not os.path.exists(path):
            print(f"  SKIP: {filename} not found")
            continue

        doc = fitz.open(path)
        page = doc[page_num - 1]
        text = page.get_text().strip()
        char_count = len(text)
        is_footer = bool(_PAGINATION_FOOTER_RE.match(text))
        effective_len = 0 if is_footer else char_count
        doc.close()

        actual = "sparse" if effective_len < 50 else "content"
        if actual == expected:
            passed += 1
        else:
            failed += 1
            print(f"  FAIL: {filename} p{page_num}: expected={expected}, got={actual} "
                  f"(chars={char_count}, footer={is_footer}, effective={effective_len})")
            print(f"    text={repr(text[:80])}")

    print(f"  {passed} passed, {failed} failed out of {len(test_cases)}")
    return failed == 0


# ── DOCX table extraction test ──

def test_docx_table_extraction():
    """Verify DOCX extraction includes table content."""
    # Create a temp DOCX with a table
    from docx import Document as DocxDocument
    import tempfile

    doc = DocxDocument()
    doc.add_paragraph("Student Application Form")
    
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Grade"
    table.cell(0, 2).text = "GPA"
    table.cell(1, 0).text = "Jane Doe"
    table.cell(1, 1).text = "11"
    table.cell(1, 2).text = "3.85"
    table.cell(2, 0).text = "John Smith"
    table.cell(2, 1).text = "10"
    table.cell(2, 2).text = "3.92"

    doc.add_paragraph("Thank you for applying.")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    try:
        text = DocumentProcessor.extract_text_from_docx(temp_path)
        print("=== DOCX Table Extraction ===")
        
        checks = [
            ("Has paragraph text", "Student Application Form" in text),
            ("Has table content", "Jane Doe" in text),
            ("Has GPA value", "3.85" in text),
            ("Has table marker", "[Table" in text),
            ("Has closing text", "Thank you for applying" in text),
        ]
        
        passed = 0
        failed = 0
        for check_name, result in checks:
            if result:
                passed += 1
            else:
                failed += 1
                print(f"  FAIL: {check_name}")
        
        print(f"  {passed} passed, {failed} failed")
        return failed == 0
    finally:
        os.unlink(temp_path)


# ── OCR triggering logic test ──

def test_ocr_triggering():
    """Verify OCR is triggered for pagination-only pages."""
    print("=== OCR Triggering Logic ===")
    
    test_pages = [
        # (text, has_images, should_trigger_ocr)
        ("Thai, Brandon - #1807\n7 of 13", True, True),      # Footer + image → OCR
        ("Thai, Brandon - #1807\n7 of 13", False, True),      # Footer only → still OCR (< 20 effective)
        ("", True, True),                                       # Empty + image → OCR
        ("", False, True),                                      # Empty, no image → OCR (< 20 chars)
        ("Short text", True, True),                             # < 20 + image → OCR
        ("A" * 200, True, False),                               # Plenty of text → no OCR
        ("Real application text about my interest in STEM and why I want to participate", False, False),  # 80+ chars → no OCR
    ]
    
    passed = 0
    failed = 0
    
    for text, has_images, expected_trigger in test_pages:
        is_pagination = bool(_PAGINATION_FOOTER_RE.match(text)) if text else False
        effective_len = 0 if is_pagination else len(text)
        
        would_trigger = (
            effective_len < 100  # Below threshold
            and (has_images or effective_len < 20 or is_pagination)  # Has reason to OCR
        )
        
        if would_trigger == expected_trigger:
            passed += 1
        else:
            failed += 1
            print(f"  FAIL: text={repr(text[:40])}, images={has_images}, "
                  f"expected={expected_trigger}, got={would_trigger} "
                  f"(effective={effective_len}, pagination={is_pagination})")
    
    print(f"  {passed} passed, {failed} failed")
    return failed == 0


if __name__ == '__main__':
    results = []
    results.append(("Pagination Footer Regex", test_pagination_footer_regex()))
    results.append(("OCR Triggering Logic", test_ocr_triggering()))
    results.append(("DOCX Table Extraction", test_docx_table_extraction()))
    results.append(("PDF Sparse Page Detection", test_pdf_extraction_stats()))
    
    print("\n=== SUMMARY ===")
    all_pass = True
    for name, passed in results:
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"  {status}: {name}")
        if not passed:
            all_pass = False
    
    sys.exit(0 if all_pass else 1)
